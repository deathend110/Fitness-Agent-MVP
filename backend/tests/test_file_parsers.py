from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from openpyxl import Workbook
from PIL import Image

from backend.files.parsers import parse_uploaded_file


def test_markdown_parser_extracts_title_text_and_line_count(tmp_path: Path) -> None:
    file_path = tmp_path / "训练说明.md"
    file_path.write_text("# 周一训练\n\n- 深蹲 5x5\n- 卧推 4x6\n", encoding="utf-8")

    result = parse_uploaded_file(file_path, ".md", "text/markdown")

    assert result["status"] == "parsed"
    assert result["summary"]["kind"] == "markdown"
    assert result["summary"]["title"] == "周一训练"
    assert "深蹲 5x5" in result["summary"]["text"]
    assert result["summary"]["preview"]["lineCount"] == 3


def test_docx_parser_reads_paragraphs_and_marks_empty(tmp_path: Path) -> None:
    docx_path = tmp_path / "coach.docx"
    document = Document()
    document.add_heading("训练反馈", level=1)
    document.add_paragraph("硬拉最后一组 RPE 9。")
    document.save(docx_path)

    parsed = parse_uploaded_file(docx_path, ".docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    assert parsed["status"] == "parsed"
    assert parsed["summary"]["kind"] == "docx"
    assert "训练反馈" in parsed["summary"]["text"]
    assert "RPE 9" in parsed["summary"]["summary"]

    empty_path = tmp_path / "empty.docx"
    Document().save(empty_path)
    empty = parse_uploaded_file(empty_path, ".docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    assert empty["status"] == "empty"
    assert empty["summary"]["text"] == ""


def test_excel_parser_summarizes_sheets_headers_and_training_columns(tmp_path: Path) -> None:
    xlsx_path = tmp_path / "log.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "June"
    sheet.append(["date", "exercise", "sets", "reps", "rpe", "weight"])
    sheet.append(["2026-06-01", "Squat", 5, 5, 8, 120])
    workbook.save(xlsx_path)

    result = parse_uploaded_file(xlsx_path, ".xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    assert result["status"] == "parsed"
    assert result["summary"]["kind"] == "excel"
    assert result["summary"]["preview"]["sheets"][0]["name"] == "June"
    assert result["summary"]["preview"]["trainingColumns"] == ["date", "exercise", "sets", "reps", "rpe", "weight"]
    assert "Squat" in result["summary"]["text"]


def test_image_parser_returns_metadata_without_embedding_full_base64(tmp_path: Path) -> None:
    image_path = tmp_path / "form.png"
    image = Image.new("RGB", (32, 24), color=(20, 40, 60))
    buffer = BytesIO()
    image.save(buffer, format="PNG")
    image_path.write_bytes(buffer.getvalue())

    result = parse_uploaded_file(image_path, ".png", "image/png")

    assert result["status"] == "parsed"
    assert result["summary"]["kind"] == "image"
    assert result["summary"]["preview"]["width"] == 32
    assert result["summary"]["preview"]["height"] == 24
    assert "base64" not in result["summary"]
